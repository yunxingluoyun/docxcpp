#!/usr/bin/env python3

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_run_font_slots(run, *, east_asia: str | None = None, complex_script: str | None = None) -> None:
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    if east_asia is not None:
        r_fonts.set(qn("w:eastAsia"), east_asia)
    if complex_script is not None:
        r_fonts.set(qn("w:cs"), complex_script)


def configure_header_footer(section, is_second: bool = False) -> None:
    if is_second:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

    header0 = section.header.paragraphs[0]
    header0.text = "Second Header" if is_second else "Py Default Header"
    header0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not is_second:
        header1 = section.header.add_paragraph("Header Tail")
        header1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer0 = section.footer.paragraphs[0]
    footer0.text = "Second Footer" if is_second else "Py Default Footer"
    footer0.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def main() -> int:
    if len(sys.argv) != 2:
        print("usage: generate_python_docx_fixture.py <output-path>", file=sys.stderr)
        return 2

    output_path = Path(sys.argv[1])
    output_path.parent.mkdir(parents=True, exist_ok=True)

    repo_root = Path(__file__).resolve().parents[1]
    image_path = repo_root / "cpp" / "image.jpeg"

    doc = Document()
    section0 = doc.sections[0]
    section0.orientation = WD_ORIENT.LANDSCAPE
    section0.page_width = Pt(720)
    section0.page_height = Pt(540)
    section0.top_margin = Pt(72)
    section0.right_margin = Pt(54)
    section0.bottom_margin = Pt(72)
    section0.left_margin = Pt(54)
    section0.header_distance = Pt(30)
    section0.footer_distance = Pt(24)
    section0.gutter = Pt(12)
    doc.settings.odd_and_even_pages_header_footer = True
    section0.different_first_page_header_footer = True

    configure_header_footer(section0)
    first_header = section0.first_page_header.paragraphs[0]
    first_header.text = "Py First Header"
    first_footer = section0.first_page_footer.paragraphs[0]
    first_footer.text = "Py First Footer"
    even_header = section0.even_page_header.paragraphs[0]
    even_header.text = "Py Even Header"
    even_footer = section0.even_page_footer.paragraphs[0]
    even_footer.text = "Py Even Footer"

    title = doc.add_heading("PythonDocx Title", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formatted = doc.add_paragraph()
    formatted.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = formatted.paragraph_format
    fmt.left_indent = Pt(24)
    fmt.right_indent = Pt(18)
    fmt.first_line_indent = Pt(-12)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(10)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(18)
    fmt.keep_together = True
    fmt.keep_with_next = True
    fmt.page_break_before = True
    fmt.widow_control = False
    fmt.tab_stops.add_tab_stop(Pt(72))
    fmt.tab_stops.add_tab_stop(Pt(144))

    run0 = formatted.add_run("Styled ")
    run0.style = "Strong"
    run0.bold = True
    run0.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    run0.font.size = Pt(16)
    run0.font.highlight_color = WD_COLOR_INDEX.YELLOW

    run1 = formatted.add_run("body")
    run1.italic = True
    run1.font.name = "Courier New"
    run1.font.size = Pt(14)
    run1.font.color.rgb = RGBColor(0xAA, 0x22, 0x11)
    run1.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    run1.add_break()
    run1.add_text("line 2")

    font_flags = doc.add_paragraph()
    caps_run = font_flags.add_run("Caps")
    caps_run.font.name = "Arial"
    caps_run.font.all_caps = True
    caps_run.font.strike = True
    set_run_font_slots(caps_run, east_asia="SimSun", complex_script="Amiri")

    sub_run = font_flags.add_run("Sub")
    sub_run.font.small_caps = True
    sub_run.font.subscript = True

    super_run = font_flags.add_run("Super")
    super_run.font.double_strike = True
    super_run.font.superscript = True

    tabs = doc.add_paragraph("Alpha\tBeta\tGamma")
    tabs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabs.paragraph_format.tab_stops.add_tab_stop(Pt(72))
    tabs.paragraph_format.tab_stops.add_tab_stop(Pt(144))

    link_paragraph = doc.add_paragraph()
    add_hyperlink(link_paragraph, "Python Link", "https://example.com/python-link")

    commented = doc.add_paragraph()
    commented_run = commented.add_run("Commented paragraph")
    doc.add_comment(commented_run, "python fixture note", "PyDocx", "PD")

    picture_paragraph = doc.add_paragraph("Body picture: ")
    picture_paragraph.add_run().add_picture(str(image_path), width=Pt(64), height=Pt(32))

    page_break_paragraph = doc.add_paragraph()
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)

    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Shading Accent 1"
    merged_head = table.cell(0, 0).merge(table.cell(0, 1))
    merged_head.text = "Merged Head"
    table.cell(0, 2).text = "H3"
    table.cell(1, 0).text = "R2C1"
    cell_with_nested = table.cell(1, 1)
    cell_with_nested.text = "Outer Cell"
    cell_link = table.cell(1, 1).add_paragraph()
    add_hyperlink(cell_link, "Cell Link", "https://example.com/cell-link")
    cell_with_picture = table.cell(1, 1).add_paragraph()
    cell_with_picture.add_run().add_picture(str(image_path), width=Pt(24), height=Pt(24))
    vertical_merge = table.cell(1, 2).merge(table.cell(2, 2))
    vertical_merge.text = "Vertical Merge"
    table.cell(2, 0).text = "R3C1"
    nested_holder = table.cell(2, 1)
    nested_holder.text = "Nested Host"
    nested = nested_holder.add_table(rows=1, cols=2)
    nested.style = "Table Grid"
    nested.cell(0, 0).text = "Nested A"
    nested.cell(0, 1).text = "Nested B"

    section1 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section1.orientation = WD_ORIENT.PORTRAIT
    section1.page_width = Pt(612)
    section1.page_height = Pt(792)
    section1.top_margin = Pt(90)
    section1.right_margin = Pt(36)
    section1.bottom_margin = Pt(54)
    section1.left_margin = Pt(36)
    section1.header_distance = Pt(18)
    section1.footer_distance = Pt(18)
    section1.gutter = Pt(0)
    section1.different_first_page_header_footer = False

    configure_header_footer(section1, is_second=True)
    doc.add_paragraph("Second section paragraph")

    doc.save(output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
