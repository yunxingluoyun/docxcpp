#!/usr/bin/env python3

from __future__ import annotations

import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT / "src"))

from docx import Document  # noqa: E402
from docx.enum.section import WD_ORIENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def comment_list(document: Document):
    return list(document.comments)


def paragraph_by_text(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise AssertionError(f"paragraph not found: {text}")


def find_paragraph_prefix(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise AssertionError(f"paragraph prefix not found: {prefix}")


def validate_all_features(path: Path) -> None:
    require(path.exists(), f"missing docx: {path}")
    doc = Document(path)

    require(len(doc.paragraphs) == 11, "all-features paragraphs mismatch")
    require(len(doc.tables) == 1, "all-features tables mismatch")
    require(len(doc.inline_shapes) == 2, "all-features images mismatch")
    require(len(doc.sections) == 2, "all-features sections mismatch")
    require(len(comment_list(doc)) == 1, "all-features comment count mismatch")

    require(doc.paragraphs[0].text == "Document Title", "all-features title mismatch")
    require(doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER, "all-features title align")
    require(doc.paragraphs[0].runs[0].style.name == "Emphasis", "all-features title character style")
    require(doc.paragraphs[0].runs[0].font.all_caps, "all-features title all caps")
    require(doc.paragraphs[5].hyperlinks[0].url == "https://openai.com", "all-features hyperlink url")

    preface = doc.paragraphs[1]
    require(preface.text == "preface note", "all-features inserted paragraph text mismatch")
    require(preface.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY, "all-features inserted paragraph align")
    require(preface.runs[0].bold, "all-features inserted paragraph run0 bold")
    require(preface.runs[0].style.name == "Strong", "all-features inserted paragraph run0 style")
    require(preface.runs[0].font.strike, "all-features inserted paragraph run0 strike")
    require(preface.runs[1].font.name == "Times New Roman", "all-features inserted paragraph run1 font")
    require(preface.runs[1].font.small_caps, "all-features inserted paragraph run1 small caps")
    require(preface.runs[2].font.superscript, "all-features inserted paragraph run2 superscript")
    require(preface.runs[2].text == " note", "all-features inserted paragraph run2 text")

    alpha = doc.paragraphs[2]
    require(alpha.text == "alpha body", "all-features alpha paragraph mismatch")
    require(round(alpha.paragraph_format.left_indent.pt) == 24, "all-features alpha left indent")
    require(alpha.paragraph_format.widow_control is False, "all-features widow control mismatch")

    comment = comment_list(doc)[0]
    require(comment.text == "alpha note", "all-features comment text mismatch")
    require(comment.author == "Codex", "all-features comment author mismatch")
    require(comment.initials == "CX", "all-features comment initials mismatch")

    require(doc.settings.odd_and_even_pages_header_footer, "all-features even/odd setting mismatch")

    section0 = doc.sections[0]
    section1 = doc.sections[1]
    require(section0.orientation == WD_ORIENT.LANDSCAPE, "all-features section0 orientation")
    require(section1.orientation == WD_ORIENT.PORTRAIT, "all-features section1 orientation")
    require(round(section0.page_width.pt) == 595, "all-features section0 width")
    require(round(section0.page_height.pt) == 842, "all-features section0 height")
    require(round(section1.page_width.pt) == 612, "all-features section1 width")
    require(round(section1.page_height.pt) == 792, "all-features section1 height")
    require(round(section0.top_margin.pt) == 72, "all-features section0 top margin")
    require(round(section0.right_margin.pt) == 54, "all-features section0 right margin")
    require(round(section0.header_distance.pt) == 36, "all-features section0 header distance")
    require(round(section0.footer_distance.pt) == 24, "all-features section0 footer distance")
    require(round(section1.top_margin.pt) == 90, "all-features section1 top margin")
    require(round(section1.left_margin.pt) == 36, "all-features section1 left margin")
    require(round(section1.header_distance.pt) == 18, "all-features section1 header distance")
    require(round(section1.footer_distance.pt) == 18, "all-features section1 footer distance")
    require(section0.different_first_page_header_footer, "all-features first-page setting")

    require([p.text for p in section0.header.paragraphs] == ["First Header", "First Header Tail", "Styled Header"],
            "all-features section0 header paragraphs")
    require([p.text for p in section0.first_page_header.paragraphs] == ["First Page Header"],
            "all-features section0 first header paragraphs")
    require([p.text for p in section1.header.paragraphs] == ["Second Header"],
            "all-features section1 header paragraphs")
    require([p.text for p in section1.even_page_header.paragraphs] == ["Even Header"],
            "all-features section1 even header paragraphs")
    require([p.text for p in section0.footer.paragraphs] == ["First Footer"],
            "all-features section0 footer paragraphs")
    require([p.text for p in section0.first_page_footer.paragraphs] == ["First Page Footer"],
            "all-features section0 first footer paragraphs")
    require([p.text for p in section1.footer.paragraphs] == ["Second Footer", "Footer Tail", "foot-mix"],
            "all-features section1 footer paragraphs")
    require([p.text for p in section1.even_page_footer.paragraphs] == ["Even Footer"],
            "all-features section1 even footer paragraphs")
    require(section1.footer.paragraphs[1].runs[0].italic, "all-features styled footer italic")
    require(section1.footer.paragraphs[1].runs[0].font.double_strike,
            "all-features styled footer double strike")
    require(str(section1.footer.paragraphs[1].runs[0].font.color.rgb) == "AA44AA",
            "all-features styled footer color")
    require(doc.paragraphs[8].runs[0].font.subscript, "all-features aside subscript")

    table = doc.tables[0]
    require(table.style.name == "Medium Grid 1 Accent 1", "all-features table style mismatch")
    require(len(table.rows) == 3, "all-features table row count mismatch")
    require(len(table.rows[0].cells) == 4, "all-features logical col count mismatch")
    require(table.cell(0, 0).text == "r1c1\nmerged alias paragraph", "all-features merged cell text")
    require(table.cell(0, 1).text == "r1c1\nmerged alias paragraph", "all-features merged alias text")
    require(table.cell(1, 2).text == "r2c3\nsecond paragraph\n", "all-features rich cell text")
    require(table.cell(1, 3).text == "r2c4", "all-features appended col text")
    require(table.cell(2, 1).text == "r3c2", "all-features third row text")
    require(len(table.cell(2, 1).tables) == 1, "all-features nested table count")
    nested = table.cell(2, 1).tables[0]
    require(len(nested.rows) == 1, "all-features nested row count")
    require(nested.cell(0, 0).text == "nested-a", "all-features nested cell 00")
    require(nested.cell(0, 1).text == "nested-b", "all-features nested cell 01")

    with zipfile.ZipFile(path) as archive:
        document_xml_text = archive.read("word/document.xml").decode("utf-8")
    require("w:eastAsia=\"Microsoft YaHei\"" in document_xml_text, "all-features missing eastAsia font slot")
    require("w:cs=\"Amiri\"" in document_xml_text, "all-features missing complex script font slot")


def validate_api_coverage(path: Path) -> None:
    require(path.exists(), f"missing docx: {path}")
    doc = Document(path)

    require(len(doc.paragraphs) == 7, "api-coverage paragraphs mismatch")
    require(len(doc.tables) == 2, "api-coverage tables mismatch")
    require(len(doc.inline_shapes) == 2, "api-coverage images mismatch")
    require(len(doc.sections) == 1, "api-coverage sections mismatch")
    require(len(comment_list(doc)) == 1, "api-coverage comment count mismatch")

    require(doc.paragraphs[4].text == "Mutable tail", "api-coverage mutable paragraph mismatch")
    require(doc.paragraphs[5].hyperlinks[0].url == "https://example.com/project",
            "api-coverage body hyperlink mismatch")

    comment = comment_list(doc)[0]
    require(comment.text == "coverage note", "api-coverage comment text mismatch")
    require(comment.author == "DocxCPP", "api-coverage comment author mismatch")
    require(comment.initials == "DC", "api-coverage comment initials mismatch")

    section = doc.sections[0]
    require(section.orientation == WD_ORIENT.LANDSCAPE, "api-coverage orientation mismatch")
    require(round(section.page_width.pt) == 612, "api-coverage width mismatch")
    require(round(section.page_height.pt) == 792, "api-coverage height mismatch")
    require(round(section.left_margin.pt) == 90, "api-coverage left margin mismatch")
    require(round(section.right_margin.pt) == 90, "api-coverage right margin mismatch")
    require(round(section.header_distance.pt) == 36, "api-coverage header distance mismatch")
    require(round(section.footer_distance.pt) == 36, "api-coverage footer distance mismatch")

    table0 = doc.tables[0]
    require(len(table0.rows) == 3, "api-coverage table0 rows mismatch")
    require(len(table0.rows[0].cells) == 3, "api-coverage table0 cols mismatch")
    require(table0.cell(0, 0).text == "A1\nA1 second", "api-coverage table0 cell00 mismatch")
    require(table0.cell(1, 0).text == "\nHeading Wrapper\nCellLink", "api-coverage table0 cell10 mismatch")
    require(table0.cell(2, 1).text == "C2\nC2 second", "api-coverage table0 cell21 mismatch")
    require(len(table0.cell(2, 2).tables) == 1, "api-coverage nested table count mismatch")
    nested = table0.cell(2, 2).tables[0]
    require(len(nested.rows) == 2, "api-coverage nested rows mismatch")
    require(nested.cell(0, 0).text == "N11", "api-coverage nested cell00 mismatch")
    require(nested.cell(0, 1).text == "Heading Wrapper", "api-coverage nested cell01 mismatch")
    require(nested.cell(1, 0).text == "\nN21", "api-coverage nested cell10 mismatch")
    require(nested.cell(1, 1).text == "N22", "api-coverage nested cell11 mismatch")

    table1 = doc.tables[1]
    require(len(table1.rows) == 1, "api-coverage table1 rows mismatch")
    require(len(table1.rows[0].cells) == 2, "api-coverage table1 logical cols mismatch")
    require(table1.cell(0, 0).text == "MergedStart\nMergedAlias", "api-coverage merged cell text")
    require(table1.cell(0, 1).text == "MergedStart\nMergedAlias", "api-coverage merged alias text")


def validate_multi_section(path: Path) -> None:
    require(path.exists(), f"missing docx: {path}")
    doc = Document(path)
    require(len(doc.sections) == 2, "multi-section section count mismatch")

    section0 = doc.sections[0]
    section1 = doc.sections[1]
    require(section0.orientation == WD_ORIENT.LANDSCAPE, "multi-section section0 orientation")
    require(round(section0.page_width.pt) == 1000, "multi-section section0 width")
    require(round(section0.page_height.pt) == 500, "multi-section section0 height")
    require(round(section0.top_margin.pt) == 72, "multi-section section0 top margin")
    require(round(section0.right_margin.pt) == 36, "multi-section section0 right margin")
    require(round(section0.bottom_margin.pt) == 108, "multi-section section0 bottom margin")
    require(round(section0.left_margin.pt) == 54, "multi-section section0 left margin")
    require(round(section0.header_distance.pt) == 27, "multi-section section0 header distance")
    require(round(section0.footer_distance.pt) == 18, "multi-section section0 footer distance")
    require(round(section0.gutter.pt) == 9, "multi-section section0 gutter")

    require(section1.orientation == WD_ORIENT.LANDSCAPE, "multi-section section1 orientation")
    require(round(section1.page_width.pt) == 612, "multi-section section1 width")
    require(round(section1.page_height.pt) == 792, "multi-section section1 height")
    require(round(section1.left_margin.pt) == 90, "multi-section section1 left margin")


def validate_full_feature_example(path: Path) -> None:
    require(path.exists(), f"missing docx: {path}")
    doc = Document(path)

    require(len(doc.sections) == 1, "full-feature sections mismatch")
    require(len(doc.tables) == 1, "full-feature tables mismatch")
    require(len(doc.inline_shapes) == 4, "full-feature images mismatch")
    require(len(comment_list(doc)) == 1, "full-feature comment count mismatch")

    section = doc.sections[0]
    require(section.orientation == WD_ORIENT.LANDSCAPE, "full-feature orientation mismatch")
    require(round(section.page_width.pt) == 612, "full-feature width mismatch")
    require(round(section.page_height.pt) == 792, "full-feature height mismatch")
    require(round(section.top_margin.pt) == 72, "full-feature top margin mismatch")
    require(round(section.right_margin.pt) == 72, "full-feature right margin mismatch")
    require(round(section.bottom_margin.pt) == 72, "full-feature bottom margin mismatch")
    require(round(section.left_margin.pt) == 72, "full-feature left margin mismatch")
    require(round(section.header_distance.pt) == 30, "full-feature header distance mismatch")
    require(round(section.footer_distance.pt) == 30, "full-feature footer distance mismatch")
    require(round(section.gutter.pt) == 12, "full-feature gutter mismatch")

    title = paragraph_by_text(doc, "DocxCPP 全功能中文示例")
    require(title.alignment == WD_ALIGN_PARAGRAPH.CENTER, "full-feature title alignment mismatch")

    formatted = find_paragraph_prefix(doc, "这个段落用于验证缩进")
    require(formatted.alignment == WD_ALIGN_PARAGRAPH.CENTER, "full-feature formatted alignment")
    require(round(formatted.paragraph_format.left_indent.pt) == 24, "full-feature left indent")
    require(round(formatted.paragraph_format.right_indent.pt) == 18, "full-feature right indent")
    require(round(formatted.paragraph_format.first_line_indent.pt) == -12, "full-feature first-line indent")
    require(round(formatted.paragraph_format.space_before.pt) == 10, "full-feature space before")
    require(round(formatted.paragraph_format.space_after.pt) == 14, "full-feature space after")
    require(round(formatted.paragraph_format.line_spacing.pt) == 20, "full-feature line spacing")

    comment = comment_list(doc)[0]
    require(comment.text == "这是针对可变段落的批注", "full-feature comment text mismatch")
    require(comment.author == "DocxCPP", "full-feature comment author mismatch")
    require(comment.initials == "DX", "full-feature comment initials mismatch")

    body_link = paragraph_by_text(doc, "项目主页")
    require(body_link.hyperlinks[0].url == "https://example.com/docxcpp-demo", "full-feature body hyperlink")

    table = doc.tables[0]
    require(table.cell(0, 0).text == "合并表头", "full-feature table merged header mismatch")
    require(table.cell(1, 0).text == "这是混合 run 段落，位于第 1 页。\n单元格中的第二段文字",
            "full-feature table cell10 mismatch")
    require(table.cell(1, 1).paragraphs[1].text == "第 2 页 段落格式", "full-feature table cell11 heading")
    require(table.cell(1, 1).paragraphs[2].hyperlinks[0].url == "https://example.com/table-link",
            "full-feature table cell hyperlink mismatch")

    with zipfile.ZipFile(path) as archive:
        document_xml_bytes = archive.read("word/document.xml")
        document_xml = ET.fromstring(document_xml_bytes)
        rels_xml = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
        comments_xml = ET.fromstring(archive.read("word/comments.xml"))
    document_xml_text = document_xml_bytes.decode("utf-8")

    explicit_breaks = document_xml.findall(".//w:br[@w:type='page']", NS)
    page_break_before = document_xml.findall(".//w:pageBreakBefore", NS)
    require(len(explicit_breaks) == 3, "full-feature explicit page break count mismatch")
    require(len(page_break_before) == 1, "full-feature pageBreakBefore count mismatch")
    require(len(explicit_breaks) + len(page_break_before) + 1 == 5, "full-feature inferred page count mismatch")

    relationship_targets = [rel.attrib.get("Target", "") for rel in rels_xml.findall(".//rel:Relationship", NS)]
    require("https://example.com/docxcpp-demo" in relationship_targets, "full-feature missing body hyperlink rel")
    require("https://example.com/table-link" in relationship_targets, "full-feature missing cell hyperlink rel")
    require("这是针对可变段落的批注" in "".join(comments_xml.itertext()), "full-feature missing comment xml")


def main() -> int:
    if len(sys.argv) != 2:
      print("usage: python_docx_regression.py <artifact-dir>", file=sys.stderr)
      return 2

    artifact_dir = Path(sys.argv[1])
    validate_all_features(artifact_dir / "all-features-generated.docx")
    validate_api_coverage(artifact_dir / "api-coverage-generated.docx")
    validate_multi_section(artifact_dir / "api-coverage-multi-section.docx")
    validate_full_feature_example(artifact_dir / "full-feature-5-pages.docx")
    print(f"python-docx regression validated: {artifact_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
