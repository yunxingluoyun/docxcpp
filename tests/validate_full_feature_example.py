#!/usr/bin/env python3

from __future__ import annotations

import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def main() -> int:
    if len(sys.argv) != 2:
        print("usage: validate_full_feature_example.py <docx-path>", file=sys.stderr)
        return 2

    path = Path(sys.argv[1])
    require(path.exists(), f"missing docx file: {path}")

    doc = Document(path)

    require(len(doc.sections) == 1, "expected exactly one writable section")
    section = doc.sections[0]
    require(section.orientation == WD_ORIENT.LANDSCAPE, "section orientation mismatch")
    require(round(section.page_width.pt) == 612, "page width mismatch")
    require(round(section.page_height.pt) == 792, "page height mismatch")
    require(round(section.top_margin.pt) == 72, "top margin mismatch")
    require(round(section.bottom_margin.pt) == 72, "bottom margin mismatch")
    require(round(section.left_margin.pt) == 72, "left margin mismatch")
    require(round(section.right_margin.pt) == 72, "right margin mismatch")
    require(round(section.header_distance.pt) == 30, "header distance mismatch")
    require(round(section.footer_distance.pt) == 30, "footer distance mismatch")
    require(round(section.gutter.pt) == 12, "gutter mismatch")

    paragraph_texts = [p.text for p in doc.paragraphs]
    require("DocxCPP 全功能中文示例" in paragraph_texts, "missing title paragraph")
    require("项目主页" in paragraph_texts, "missing hyperlink paragraph text")
    require(
        any("这个段落用于验证缩进" in text for text in paragraph_texts),
        "missing formatted paragraph",
    )
    require(
        any("第 4 页通过当前段落上的 page_break_before 生成" in text for text in paragraph_texts),
        "missing page 4 intro paragraph",
    )
    require(
        any("第 5 页用于收尾" in text for text in paragraph_texts),
        "missing closing summary paragraph",
    )

    title = next(p for p in doc.paragraphs if p.text == "DocxCPP 全功能中文示例")
    require(title.alignment == WD_ALIGN_PARAGRAPH.CENTER, "title alignment mismatch")

    formatted = next(p for p in doc.paragraphs if p.text.startswith("这个段落用于验证缩进"))
    require(formatted.alignment == WD_ALIGN_PARAGRAPH.CENTER, "formatted paragraph alignment mismatch")
    require(round(formatted.paragraph_format.left_indent.pt) == 24, "left indent mismatch")
    require(round(formatted.paragraph_format.right_indent.pt) == 18, "right indent mismatch")
    require(round(formatted.paragraph_format.first_line_indent.pt) == -12, "first line indent mismatch")
    require(round(formatted.paragraph_format.space_before.pt) == 10, "space before mismatch")
    require(round(formatted.paragraph_format.space_after.pt) == 14, "space after mismatch")
    require(round(formatted.paragraph_format.line_spacing.pt) == 20, "line spacing mismatch")

    require(len(doc.tables) == 1, "expected one table")
    table = doc.tables[0]
    require(table.cell(0, 0).text == "合并表头", "merged header cell mismatch")
    require(
        "这是混合 run 段落，位于第 1 页。\n单元格中的第二段文字" == table.cell(1, 0).text,
        "table cell paragraph content mismatch",
    )
    require("第 2 页 段落格式\n单元格超链接" in table.cell(1, 1).text, "table hyperlink cell mismatch")

    require(len(doc.inline_shapes) == 4, "expected four inline pictures")

    with zipfile.ZipFile(path) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))
        rels_xml = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
        comments_xml = ET.fromstring(archive.read("word/comments.xml"))

    explicit_breaks = document_xml.findall(".//w:br[@w:type='page']", NS)
    page_break_before = document_xml.findall(".//w:pageBreakBefore", NS)
    require(len(explicit_breaks) == 3, "expected three explicit page breaks")
    require(len(page_break_before) == 1, "expected one pageBreakBefore paragraph")
    require(len(explicit_breaks) + len(page_break_before) + 1 == 5, "expected five inferred pages")

    hyperlinks = document_xml.findall(".//w:hyperlink", NS)
    require(len(hyperlinks) == 2, "expected two hyperlink nodes")

    relationship_targets = [rel.attrib.get("Target", "") for rel in rels_xml.findall(".//rel:Relationship", NS)]
    require(
        "https://example.com/docxcpp-demo" in relationship_targets,
        "missing document hyperlink target",
    )
    require(
        "https://example.com/table-link" in relationship_targets,
        "missing table hyperlink target",
    )

    comment_text = "".join(comments_xml.itertext())
    require("这是针对可变段落的批注" in comment_text, "missing comment text")

    print(f"validated={path}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"images={len(doc.inline_shapes)}")
    print("inferred_pages=5")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
